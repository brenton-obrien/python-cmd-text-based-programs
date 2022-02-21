# TEXT BASED PYTHON PROGRAM
# A cmd based program that will complete simple text based functions such as:

# Fizz Buzz
# Reverse a string
# Pig Latin
# Count Vowels
# Check if Palindrome
# Count words in a string
# Journal
# Vigenere/Vernam/ Caesar Cipher
# Regex Query Tool.

import os
import sys
import datetime
import random
import docx
from docx import Document
import re


def cls():  # Clears the terminal screen
    os.system('cls')


def replay_prompt(program_to_run, title):  # This runs everytime a function selected from the menu is completed

    print("--------------------------------------")
    print("1) Replay")
    print("2) Return to Menu")
    print("Q) Quit")
    print("--------------------------------------")
    print("Please enter a number to make a selection: ")
    replay_selection = input("> ")

    if replay_selection == '1':
        cls()
        program_to_run()

    elif replay_selection == '2':
        cls()
        main_menu()

    elif replay_selection.lower() == 'q':
        sys.exit()

    else:
        cls()
        print("--------------------------------------")
        print(f"{title} - Brenton O'Brien")

        replay_prompt(program_to_run, title)


def main_menu():  # Accepts user input and will run the desired program

    cls()

    print("--------------------------------------")
    print("TEXT BASED PROGRAMS - Brenton O'Brien")
    print("--------------------------------------")

    print("\n1) FizzBuzz")
    print("2) Reverse a string")
    print("3) Pig Latin")
    print("4) Count Vowels")
    print("5) Check if Palindrome")
    print("6) Count words in a string")
    print("7) Journal")
    print("8) Vingenere/ Vernam/ Casesar Cipher")
    print("9) Regex Query Tool")
    print("Q) Quit")
    print("\n--------------------------------------")
    print("Please enter a number to make a selection: ")
    menu_selection = input("> ")

    cls()

    if menu_selection == '1':
        fizz_buzz()

    elif menu_selection == '2':
        reverse_a_string()

    elif menu_selection == '3':
        pig_latin()

    elif menu_selection == '4':
        count_vowels()

    elif menu_selection == '5':
        check_if_palindrome()

    elif menu_selection == '6':
        number_of_words_in_string()

    elif menu_selection == '7':
        journal()

    elif menu_selection == '8':
        ciphers()

    elif menu_selection == '9':
        regex_query()

    elif menu_selection.lower() == 'q':
        sys.exit()

    else:
        main_menu()


def fizz_buzz():  # Completes the FizzBuzz sequence

    print("--------------------------------------")
    print("FIZZ BUZZ - Brenton O'Brien")
    print("--------------------------------------")

    try:
        fizz_buzz_range = int(input('Please enter the highest number to generate fizz buzz to:\n>'))
    except ValueError:
        cls()
        fizz_buzz()

    list_of_nums = [num for num in range(1, fizz_buzz_range + 1)]  # Creates a list of numbers from 1 to 100.

    for num in list_of_nums:  # Loops through every number and create a new print statement based on the result of the modulus checks

        # These checks are done in this specific order so that 'Fizzbuzz' numbers arnt first captured by 'Fizz' or 'Buzz' logic

        if num % 3 == 0 and num % 5 == 0:  # If number is divisible by 3 and 5
            print(f'{num}: FizzBuzz')

        elif num % 3 == 0:  # If number is divisible by 3 only
            print(f'{num}: Fizz')

        elif num % 5 == 0:  # If number is divisible by 5 only
            print(f'{num}: Buzz')

        else:  # If number is not divisible
            print(f'{num}:')

    replay_prompt(fizz_buzz, "FIZZ BUZZ")


def reverse_a_string():  # Reverses any string that is inputted

    print("--------------------------------------")
    print("REVERSE A STRING - Brenton O'Brien")
    print("--------------------------------------")

    reverse_input = input('Enter a string to be reversed:\n>')  # Get user input
    print(f"\nReversed string:\n{reverse_input[::-1]}")  # Reverse it using indexing  [<start>:<stop>:<step>]. If step is -1 it reverses the string.

    replay_prompt(reverse_a_string, "REVERSE A STRING")


def pig_latin():  # converts any string into pig latin

    print("--------------------------------------")
    print("PIG LATIN - Brenton O'Brien")
    print("--------------------------------------")

    punctuation = '''!()-[]{};:'"\,<>./?@#$%^&*_~'''  # Used to filter punctuation out of the string
    pig_latin_lower_no_punc = ''  # Placeholder to store string with no punctuation
    pig_output = ''  # Used to print out the final result
    pig_latin_input = input('Enter a string to be converted to pig latin:\n>').lower()  # Gets the initial string and converts it to lowercase

    # Strips punctuation from string and saves it to new variable
    for char in pig_latin_input:  # For each letter in input
        if char not in punctuation:  # If the letter isn't punctuation
            pig_latin_lower_no_punc += char  # Append the letter to new string

    # Defines variables which handle the if a word starts with a vowel, a two constant cluster or three constant cluster
    vowels = ('a', 'e', 'i', 'o', 'u')
    two_constant_custer = (
        'sm', 'sn', 'th', 'st', 'sw', 'sk', 'sl', 'sl', 'dw', 'tw', 'dr', 'tr', 'qu', 'cr', 'cl', 'pr', 'fr', 'br', 'gr', 'pl', 'fl', 'bl', 'gl')
    three_constant_custer = ('sph', 'thw', 'thr', 'shr')

    for word in pig_latin_lower_no_punc.split():  # Spits each word up into a list so that the following logic can loop through each word

        if word.startswith(vowels):  # If word starts with a vowel then just add '-ay' to the end of the string
            pig_output += word + '-ay '  # Add the pig latin word to the final string

        elif word.startswith(three_constant_custer):  # If word starts with a three constant cluster
            pig_output += word[3::] + '-' + word[0] + word[1] + word[2] + 'ay '  # Move the three letter to the end of the word and add '-ay'

        elif word.startswith(two_constant_custer):  # If word starts with a two constant cluster
            pig_output += word[2::] + '-' + word[0] + word[1] + 'ay '  # Move the two letter to the end of the word and add '-ay'

        else:  # If the word doesn't fit any other criteria then is must just start with a constant
            pig_output += word[1::] + '-' + word[0] + 'ay '  # Move the first letter to the end of the word and add '-ay'

    print(f"\nPig Latin:\n{pig_output}")  # Prints the completed string to the user

    replay_prompt(pig_latin, "PIG LATIN")


def count_vowels():  # Counts the vowels found in a string

    print("--------------------------------------")
    print("COUNT VOWELS - Brenton O'Brien")
    print("--------------------------------------")

    vowel_counting_input = input('Enter a sentence to have the total vowels counted:\n>')  # Input string
    print(f"\nNumber of A's: {vowel_counting_input.count('A') + vowel_counting_input.count('a')}\n"  # Adds uppercase and lowercase vowels together
          f"Number of E's: {vowel_counting_input.count('E') + vowel_counting_input.count('e')}\n"
          f"Number of I's: {vowel_counting_input.count('I') + vowel_counting_input.count('i')}\n"
          f"Number of O's: {vowel_counting_input.count('O') + vowel_counting_input.count('o')}\n"
          f"Number of U's: {vowel_counting_input.count('u') + vowel_counting_input.count('u')}\n")

    replay_prompt(count_vowels, "COUNT VOWELS")


def check_if_palindrome():  # Input a string to see if it is a palindrome

    print("--------------------------------------")
    print("CHECK IF PALINDROME - Brenton O'Brien")
    print("--------------------------------------")

    punctuation = '''!()-[]{};:'"\,<>./?@#$%^&*_~'''  # Used to filter punctuation out of the string
    palindrome_comparison = ''

    palindrome_input = input('Enter a sentence to check if palindrome:\n>')  # Gets a string from the user, converts it to all lowercase

    for char in palindrome_input.lower():
        if char in punctuation:
            pass
        else:
            palindrome_comparison += char

    palindrome_comparison = palindrome_comparison.replace(' ', '')

    if palindrome_comparison == palindrome_comparison[::-1]:  # Check if string is equal to reversed version of string
        print(f'\n"{palindrome_input}" is a palindrome!')
    else:
        print(f'\n"{palindrome_input}" is not a palindrome!')

    replay_prompt(check_if_palindrome, "CHECK IF PALINDROME")


def number_of_words_in_string():
    print("--------------------------------------")
    print("COUNT THE WORDS - Brenton O'Brien")
    print("--------------------------------------")

    word_count_input = input('Enter a sentence to count the number of words:\n>')
    print(f"\nWord Count: {len(word_count_input.split())}")

    replay_prompt(number_of_words_in_string, "COUNT THE WORDS")


# Journal - A simple application that allows people to write journal entries to word docs (Open/edit/save word docs).
def journal():
    # Function scoped variables
    document = Document()
    current_working_directory = os.getcwd()

    # Prints out the menu and awaits user input
    print("--------------------------------------")
    print("JOURNAL - Brenton O'Brien")
    print("--------------------------------------")

    print("\n1) Make a new entry")
    print("2) Edit an existing entry")
    print("3) View list of entries")
    print("4) View Specific entry")
    print("5) Delete an entry")
    print("6) Return to menu")
    print("Q) Quit")

    print("\n--------------------------------------")
    journal_selection = input("Please enter a number to make a selection:\n>")

    # Menu logic
    # 1) Makes a new entry
    if journal_selection == '1':

        # Gets a time stamp that can be used as the filename
        current_time = datetime.datetime.now().strftime("%d %b %y - %I-%M-%S %p")  #

        cls()  # Clear screen and reprint title
        print("--------------------------------------")
        print("JOURNAL - Brenton O'Brien")
        print("--------------------------------------")

        # Gets the users text
        user_text = input(f'{current_time}\n\nEnter text, press <ENTER> key to submit entry:\n--------------------------------------\n\n>')

        # Adds it to a document
        document.add_paragraph(user_text)

        # Saves that document (time stamp is the filename)
        document.save(f'diary_entries/{datetime.datetime.now().strftime("%d %b %y - %I-%M-%S %p")}.docx')

        print("\n--------------------------------------")

        # Inform user that the entry has been saved
        print(f'\nEntry saved as "diary_entries/{current_time}.docx"\n')

        # Prompt replay
        replay_prompt(journal, "JOURNAL")

    # 2) Edit existing Entry
    elif journal_selection == '2':

        file_found = False

        cls()  # Clears screen and reprint title
        print("--------------------------------------")
        print("JOURNAL - Brenton O'Brien")
        print("--------------------------------------")

        # Asks the requested filename from the user
        edit_journal_search = input('\nEnter the exact filename of the entry you wish to edit:\n>')

        # Create a list of all current entries
        all_entries = os.listdir(f'{current_working_directory}\\diary_entries')

        print("\n--------------------------------------")

        # Loops through the diary_entries folder to see if the requested file is in there
        for file in all_entries:

            if file == edit_journal_search:  # If it is there
                print(f'\n"{file}" has been found.\n')  # Let the user know
                file_found = True  # Sets this to true so that later code is skipped
                document_text = ''  # Preps an empty string so that all the original text paragraphs can be concatenated to it

                # Creates a document object using the filepath (this is only used to 'read' the old file)
                doc = docx.Document(f"{current_working_directory}\\diary_entries\\{file}")

                # Adds all paragraphs to the empty string
                for para in doc.paragraphs:
                    document_text += para.text

                # Show this text to the user
                print(f"Original Text: \n{document_text}")

                # Delete the old file
                os.remove(f"{current_working_directory}\\diary_entries\\{file}")

                print("\n--------------------------------------")

                # Get replacement text from the user
                document.add_paragraph(input("Enter the new text:\n>"))

                # Create a new file with the same exact name
                document.save(f'diary_entries/{file}')

                print("\n--------------------------------------")

                # Let the user know that the edit has been successful
                print(f'\n{file} has been successfully edited\n')

        # If the file cannot be found, let the user know
        if not file_found:
            print(f'\nCould not find: {edit_journal_search}\n')

        # Prompt the replay regardless of whether the file was found or not
        replay_prompt(journal, "JOURNAL")

    # 3) View list of diary entries
    elif journal_selection == '3':

        # Clears the screen, reprint title and current diary directory
        cls()
        print("--------------------------------------")
        print("JOURNAL - Brenton O'Brien")
        print("--------------------------------------")
        print(f'Viewing all diary entries in: "{current_working_directory}"')
        print("--------------------------------------\n")

        # Create a list of all the entries
        all_entries = os.listdir(f'{current_working_directory}\\diary_entries')

        # Let the user know if the directory is empty
        if len(all_entries) == 0:
            print('No files found.')

        # Otherwise, show the user what files were found
        else:
            for file in all_entries:
                print(file)

        # Used as spacing
        print('')

        # Replay prompt
        replay_prompt(journal, "JOURNAL")

    # 4) View specific entry
    elif journal_selection == '4':

        file_found = False

        # Clears the screen and reprints title
        cls()
        print("--------------------------------------")
        print("JOURNAL - Brenton O'Brien")
        print("--------------------------------------")

        # Asks for the filename to be edited
        edit_journal_search = input('\nEnter the exact filename of the entry you wish to view:\n>')

        # Creates a list of files in the directory
        all_entries = os.listdir(f'{current_working_directory}\\diary_entries')

        print("\n--------------------------------------")

        # Loop through all the files and enact the following code if a match is found
        for file in all_entries:
            if file == edit_journal_search:
                file_found = True

                # Used as spacing
                print('')

                # Create document object (this is only used to print the original text from the document)
                doc = docx.Document(f"{current_working_directory}\\diary_entries\\{file}")

                # Print every paragraph to the screen
                for para in doc.paragraphs:
                    print(para.text)

                # Spacing
                print('')

        # If there is not a match, let the user know
        if not file_found:
            print(f'\nCould not find: {edit_journal_search}\n')

        # Prompt replay
        replay_prompt(journal, "JOURNAL")

    # Delete entry
    elif journal_selection == '5':

        file_found = False

        # Clear screen and reprint title
        cls()
        print("--------------------------------------")
        print("JOURNAL - Brenton O'Brien")
        print("--------------------------------------")

        # Asks for filename to be deleted
        edit_journal_search = input('\nEnter the exact filename of the entry you wish to delete:\n>')

        # Creates a list of files in the directory
        all_entries = os.listdir(f'{current_working_directory}\\diary_entries')

        print("\n--------------------------------------")

        # Loops through all files
        for file in all_entries:
            if file == edit_journal_search:
                file_found = True

                # Deletes the file and informs the user
                os.remove(f"{current_working_directory}\\diary_entries\\{file}")
                print(f"\n{file} has been deleted\n")

        # If the file couldn't be found then inform the user
        if not file_found:
            print(f'\nCould not find: {edit_journal_search}\n')

        # Prompt replay
        replay_prompt(journal, "JOURNAL")

    # Return to menu
    elif journal_selection == '6':
        cls()
        main_menu()

    # Quits
    elif journal_selection.lower() == 'q':
        quit()

    # Any other clears the screen and reruns the journal menu
    else:
        cls()
        journal()


# Vigenere/ Vernam/ Caesar Cipher
def ciphers():
    # Prints out the menu and awaits user input
    print("--------------------------------------")
    print("CIPHER's - Brenton O'Brien")
    print("--------------------------------------")

    print("\n1) Encode message with Caesar Cipher")
    print("2) Encode message with Vigenere Cipher")
    print("3) Encode message with Vernam Cipher")
    print("4) Return to menu")
    print("Q) Quit")

    print("\n--------------------------------------")
    cipher_selection = input("Please enter a number to make a selection:\n>")

    # Menu logic
    # 1) Caesar Cipher
    if cipher_selection == '1':
        cls()
        caesar_cipher()

    elif cipher_selection == '2':
        cls()
        vigenere_cipher()

    elif cipher_selection == '3':
        cls()
        vernam_cipher()

    elif cipher_selection == '4':
        cls()
        main_menu()

    elif cipher_selection.lower() == 'q':
        quit()

    else:
        cls()
        ciphers()


# CAESAR CIPHER
def caesar_cipher():
    cls()  # Clear screen and reprint title
    print("--------------------------------------")
    print("CAESAR CIPHER - Brenton O'Brien")
    print("--------------------------------------")

    # Creates a list of all letters
    alphabet_list = 'abcdefghijklmnopqrstuvwxyz'

    # Placeholder for finished output
    encrypted_text = ''

    # Amount of places the characters are to be shifted
    caesar_shift = 0

    # Gets a valid 'cesar_shift' input from the user
    try:
        caesar_shift = int(input('How many digits would you like to shift the text? \n(Enter a number between 1-23):\n>'))
        if caesar_shift not in range(1, 24):
            cls()
            caesar_cipher()

    except ValueError:
        cls()
        caesar_cipher()

    print("--------------------------------------")
    # Gets message that requires encoding from the user
    caesar_original_message = input('Enter your message to encode\n(numbers, spaces, and special character remain unencoded):\n\n>')

    # Converts user message to a list of characters
    list_of_characters = [character for character in caesar_original_message]

    # Loop through these characters
    for character in list_of_characters:

        # Only do the following if the character is in the alphabet (spaces and special character are skipped here)
        if character.isalpha():

            # Gets the characters index position and then adds the caesar shift to it (this variable is the index of the replacement character)
            shifted_index = alphabet_list.index(character.lower()) + caesar_shift

            # Allows the new index to wrap around the alphabet to prevent index out of scope errors
            if shifted_index >= 26:
                shifted_index -= 26

            # The following 4 lines allow character to retain their capitalisation (if any)
            if character.islower():
                encrypted_text += alphabet_list[shifted_index]

            elif character.isupper():
                encrypted_text += (alphabet_list[shifted_index]).upper()

        # If the character is not in the alphabet then no need to encode it (just add it the to encrypted text as is)
        else:
            encrypted_text += character

    print("\n--------------------------------------")

    # Show the user the output
    print(f"\nEncrypted message: {encrypted_text}\n")

    # Prompt replay
    replay_prompt(caesar_cipher, "CAESAR CIPHER")


# Vigenere Cipher
def vigenere_cipher():
    cls()  # Clear screen and reprint title
    print("--------------------------------------")
    print("VIGENERE CIPHER - Brenton O'Brien")
    print("--------------------------------------")

    # Creates a list of all letters
    alphabet_list = 'abcdefghijklmnopqrstuvwxyz'

    # Placeholder for finished output
    encrypted_text = ''

    key_index = 0

    # Gets a valid 'Vigenere_shift' input from the user
    vigenere_key = input('Enter your secret key: \n(single word with alphabetic characters only)\n>')

    if not vigenere_key.isalpha():
        # cls()
        vigenere_cipher()

    print("--------------------------------------")
    # Gets message that requires encoding from the user
    vigenere_original_message = input('Enter your message to encode\n(numbers, spaces, and special character remain unencoded):\n\n>')

    # Converts user message to a list of characters
    list_of_characters = [character for character in vigenere_original_message]

    # Creates a key stream (repeat the character of the key equal to the length of the message. i.e if key is cat, and its a long message, the keystream would be catcatcatcat...etc)
    key_index = 0
    key_stream = ''
    for character in list_of_characters:
        if character.isalpha():
            try:
                key_stream += (vigenere_key[key_index])
                key_index += 1
            except IndexError:
                key_index = 0
                key_stream += (vigenere_key[key_index])
                key_index += 1

    # Reset key_index so that it can be reused
    key_index = 0

    # Loop through these characters (using the index position of the next character in the key stream for each one)
    for character in list_of_characters:

        # Reset this on each iteration of the loop
        shifted_index = 0

        # Only do the following if the character is in the alphabet (spaces and special character are skipped here)
        if character.isalpha():

            # Gets the characters index position and then adds the index position of the next character in the key stream
            shifted_index = alphabet_list.index(character.lower()) + alphabet_list.index(key_stream[key_index])

            # Allows the new index to wrap around the alphabet to prevent index out of scope errors
            while shifted_index >= 26:
                shifted_index -= 26

            # The following 4 lines allow character to retain their capitalisation (if any)
            if character.islower():
                encrypted_text += alphabet_list[shifted_index]

            elif character.isupper():
                encrypted_text += (alphabet_list[shifted_index]).upper()

            try:
                key_index += 1
            except IndexError:
                key_index = 0

        # If the character is not in the alphabet then no need to encode it (just add it the to encrypted text as is)
        else:
            encrypted_text += character

    print(key_stream)

    print("\n--------------------------------------")

    # Show the user the output
    print(f"\nEncrypted message: {encrypted_text}\n")

    # Prompt replay
    replay_prompt(vigenere_cipher, "VIGENERE CIPHER")


# Vernam Cipher
def vernam_cipher():
    cls()  # Clear screen and reprint title
    print("--------------------------------------")
    print("VERNAM CIPHER - Brenton O'Brien")
    print("--------------------------------------")

    # Creates a list of all letters
    alphabet_list = 'abcdefghijklmnopqrstuvwxyz'

    # Placeholder for finished output
    encrypted_text = ''

    # Gets message that requires encoding from the user
    vernam_original_message = input('Enter your message to encode\n(numbers, spaces, and special character remain unencoded):\n\n>')

    print("--------------------------------------")

    # Converts user message to a list of characters
    list_of_characters = [character for character in vernam_original_message]

    # Creates a key stream (a random string of characters of equal length to the alphabetical characters in the original message)
    key_stream = ''
    for character in list_of_characters:
        if character.isalpha():
            key_stream += random.choice(alphabet_list)

    # Used to loop through the randomly generated key
    key_index = 0

    # Loop through these characters (using the index position of the next character in the key stream for each one)
    for character in list_of_characters:

        # Reset this on each iteration of the loop
        shifted_index = 0

        # Only do the following if the character is in the alphabet (spaces and special character are skipped here)
        if character.isalpha():

            # Gets the characters index position and then adds the index position of the next character in the key stream
            shifted_index = alphabet_list.index(character.lower()) + alphabet_list.index(key_stream[key_index])

            # Allows the new index to wrap around the alphabet to prevent index out of scope errors
            while shifted_index >= 26:
                shifted_index -= 26

            # The following 4 lines allow character to retain their capitalisation (if any)
            if character.islower():
                encrypted_text += alphabet_list[shifted_index]

            elif character.isupper():
                encrypted_text += (alphabet_list[shifted_index]).upper()

            try:
                key_index += 1
            except IndexError:
                key_index = 0

        # If the character is not in the alphabet then no need to encode it (just add it the to encrypted text as is)
        else:
            encrypted_text += character

    print(f"\nRandomly Generated Key: {key_stream}\n")

    print("--------------------------------------")

    # Show the user the output
    print(f"\nEncrypted message: {encrypted_text}\n")

    # Prompt replay
    replay_prompt(vernam_cipher, "VERNAM CIPHER")


# Regex Query Tool - Allows user to enter text or a .docx file and run their regex searches for strings of their own choosing
def regex_query():
    # Prints out the menu and awaits user input
    print("--------------------------------------")
    print("REGEX TOOL - Brenton O'Brien")
    print("--------------------------------------")

    print("\n1) Find all")
    print("2) Partial Match")
    print("3) Starts with")
    print("4) Ends with")
    print("5) Email Search")
    print("6) Phone Numbers")
    print("7) Return to menu")
    print("Q) Quit")

    print("\n--------------------------------------")
    regex_selection = input("Please enter a number to make a selection:\n>")

    # Menu logic
    # 1) Find all
    if regex_selection == '1':
        cls()
        regex_func("FIND ALL - REGEX", True, 'Enter the word(s) you want to search for in the above text:\n>')

    # 2) Partial Match
    elif regex_selection == '2':
        cls()
        regex_func("PARTIAL SEARCH - REGEX", True, 'Enter the word(s) you want to search for in the above text\n'
                                                   'Use a "." as a wildcard character:\n>')

    # 3) Starts with
    elif regex_selection == '3':
        cls()
        regex_func("STARTS WITH SEARCH - REGEX", True, 'Enter the word(s) to see if the text starts with it\n'
                                                       'Use a "^" before entering your word(s):\n>')

    # 4) Ends with
    elif regex_selection == '4':
        cls()
        regex_func("ENDS WITH SEARCH - REGEX", True, 'Enter the word(s) to see if the text ends with it\n'
                                                     'Use a "$" after entering your word(s):\n>')

    # 5) Email Search
    elif regex_selection == '5':
        cls()
        regex_func("EMAIL SEARCH - REGEX", False, 'Press <ENTER> to begin searching for all emails')

    # 6) Phone Numbers
    elif regex_selection == '6':
        regex_func("PHONE NUMBER SEARCH - REGEX", False, 'Press <ENTER> to begin searching for all phone numbers that begin with "04"')

    # 7) Return to menu
    elif regex_selection == '7':
        cls()
        main_menu()

    # q) QUIT
    elif regex_selection.lower() == 'q':
        quit()

    # Incorrect input handling
    else:
        cls()
        regex_query()


# All regex functions rely on this code, the variable "matches" will be changed (along with other specific code) based on the "title" of the selected regex query
def regex_func(title, ask_for_input, prompt):
    global search_string, original_string

    cls()  # Clear screen and reprint title
    print("--------------------------------------")
    print(f"{title} - Brenton O'Brien")
    print("--------------------------------------\n")

    print("1) Search user entered text")
    print('2) Search word document in the "regex_files" directory')
    print("3) Return to menu")

    print("\n--------------------------------------")

    # Gets user input from the menu screen
    regex_selection = input("Please enter a number to make a selection:\n>")

    # 1) Search user entered text
    if regex_selection == '1':

        print("--------------------------------------")

        # Get the original text from the user
        original_string = input('Enter text here: \n>')

        print("--------------------------------------")

        # This logic changes the regex in matches depending on the title
        if title in ["FIND ALL - REGEX", "PARTIAL SEARCH - REGEX", "STARTS WITH SEARCH - REGEX", "ENDS WITH SEARCH - REGEX"]:

            # Get word(s) that they want to search for in the above string
            search_string = input('Enter the word(s) you want to search for in the above text:\n>')

            matches = re.findall(search_string, original_string, flags=re.IGNORECASE)

        elif title == "EMAIL SEARCH - REGEX":
            search_string = 'Email'
            matches = re.findall(r"[\w\.-]+@[\w\.-]+", original_string)

        elif title == "PHONE NUMBER SEARCH - REGEX":
            search_string = 'Phone numbers'
            matches = re.findall(r"04[\d]+", original_string)

        # If no matches then let the user know
        if len(matches) == 0:
            print("\n--------------------------------------")
            print('0 matches found.')

        # If there were matches then let the user know how many there was
        else:
            print("--------------------------------------\n")
            print(f'{search_string} was found a total of {len(matches)} times.\n')
            for match in matches:
                print(match)

            print('')

        # Regardless of outcome ask user if they would like to replay
        replay_prompt(regex_query, title)

    elif regex_selection == '2':

        # Gets the current working directory
        cwd = os.getcwd()

        print("--------------------------------------")

        # Gets the requested filename
        filename = input('Enter the filename of the docx document found in the "regex_files" directory (including extension):\n>')

        # Create a list of all files in "regex files"
        all_regex_files = os.listdir(f'{cwd}\\regex_files')

        if filename not in all_regex_files:
            print("--------------------------------------")
            print(f'{filename} could not be found in "{cwd}"\n\nThe directory currently contains the following files:\n')

            for file in all_regex_files:
                print(file)

            print('')

            # Replay the menu that the user is currently in
            replay_prompt(regex_query, title)

        else:
            # Creates a document object of the requested name
            doc = docx.Document(f'{cwd}\\regex_files\\{filename}')

            # Placeholder string to store all text from the document
            all_text = ''

            # Append all text from the document object to 'all_text'
            for para in doc.paragraphs:
                all_text += ' ' + para.text

            print("--------------------------------------")

            if ask_for_input:
                # Get word(s) that they want to search for in the above string
                search_string = input(prompt)

                print("--------------------------------------")

            # This logic changes the regex in matches depending on the title
            if title in ["FIND ALL - REGEX", "PARTIAL SEARCH - REGEX", "STARTS WITH SEARCH - REGEX", "ENDS WITH SEARCH - REGEX"]:
                matches = re.findall(search_string, all_text, flags=re.IGNORECASE)

            elif title == "EMAIL SEARCH - REGEX":
                search_string = 'Email'
                matches = re.findall(r'[\w\.-]+@[\w\.-]+', all_text)

            elif title == "PHONE NUMBER SEARCH - REGEX":
                search_string = 'Phone numbers'
                matches = re.findall(r"04[\d]+", all_text)

            # If no matches then let the user know
            if len(matches) == 0:
                print('0 matches found.')

            # If there were matches then let the user know how many there was
            else:
                print(f'\n"{search_string}" was found a total of {len(matches)} times.\n')

                for match in matches:
                    print(match)

                print('')

            # Regardless of outcome ask user if they would like to replay
            replay_prompt(regex_query, title)

    elif regex_selection == '3':
        cls()
        regex_query()

    else:
        cls()
        regex_query()


main_menu()
